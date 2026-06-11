"""Tests for the resume parser."""
from __future__ import annotations

import io

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import NameObject, TextStringObject

from app.services.parser import UnsupportedFileType, extract_text


def test_extract_txt_returns_decoded_text():
    text = "Hello, 世界\nLine two"
    out = extract_text("resume.txt", text.encode("utf-8"))
    assert out == text


def test_extract_md_works_like_txt():
    out = extract_text("resume.md", b"# Title\n\nbody")
    assert "Title" in out


def test_extract_docx_collects_paragraphs_and_tables():
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Senior Engineer at Acme")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skill"
    table.rows[0].cells[1].text = "Python"

    buf = io.BytesIO()
    doc.save(buf)

    out = extract_text("resume.docx", buf.getvalue())
    assert "Jane Doe" in out
    assert "Senior Engineer at Acme" in out
    assert "Python" in out


def test_extract_pdf_returns_some_text():
    # Build a minimal PDF in-memory. pypdf's writer doesn't draw text natively,
    # so we just smoke-test that the parser tolerates an empty PDF without raising.
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buf = io.BytesIO()
    writer.write(buf)
    out = extract_text("resume.pdf", buf.getvalue())
    assert isinstance(out, str)  # may be empty string for blank PDF


def test_unsupported_extension_raises():
    with pytest.raises(UnsupportedFileType):
        extract_text("resume.rtf", b"junk")
