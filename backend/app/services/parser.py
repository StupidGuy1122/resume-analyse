"""Resume parsing — extract plain text from PDF / DOCX / TXT."""
from __future__ import annotations

import io
from pathlib import PurePath

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class UnsupportedFileType(ValueError):
    """Raised when the uploaded file extension is not supported."""


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch by file extension and return concatenated plain text."""
    suffix = PurePath(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")
    raise UnsupportedFileType(
        f"Unsupported file type: {suffix!r}. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
    )


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            # Fail-soft on a single bad page so the rest of the resume still loads.
            continue
    return "\n".join(p.strip() for p in pages if p and p.strip())


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Tables sometimes hold structured CV info — flatten cell-by-cell.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)
